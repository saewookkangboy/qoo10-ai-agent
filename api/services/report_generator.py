"""
ë¦¬í¬íŠ¸ ìƒì„± ì„œë¹„ìŠ¤
ë¶„ì„ ê²°ê³¼ë¥¼ PDF, Excel, Markdown, DOC í˜•íƒœë¡œ ë‹¤ìš´ë¡œë“œ ê°€ëŠ¥í•œ ë¦¬í¬íŠ¸ë¥¼ ìƒì„±í•©ë‹ˆë‹¤.

ë¦¬í¬íŠ¸ ìƒì„± ì›ì¹™:
- CRAWLING_ANALYSIS_PRINCIPLES.md ì°¸ì¡°
- ëª¨ë“  ë¦¬í¬íŠ¸ëŠ” ì¼ê´€ëœ êµ¬ì¡°ì™€ í˜•ì‹ì„ ë”°ë¼ì•¼ í•¨
- í¬ë¡¤ë§ ë°©ë²•(crawled_with)ì„ ëª…ì‹œí•´ì•¼ í•¨
- ì ìˆ˜ ê³„ì‚° ê¸°ì¤€ì€ ì›ì¹™ ë¬¸ì„œë¥¼ ì¤€ìˆ˜í•´ì•¼ í•¨
"""
from typing import Dict, Any, Optional, List
from datetime import datetime
import json
import io
import os

from services.logging_utils import log_debug as _log_debug

# XML ë³´ì•ˆ: defusedxmlì„ ì‚¬ìš©í•˜ì—¬ XML bomb/vector ê³µê²© ë°©ì§€
try:
    import defusedxml.ElementTree as ET
    import defusedxml
    defusedxml.defuse_stdlib()
except ImportError:
    import warnings
    warnings.warn(
        "defusedxml is not installed. XML parsing is not protected against "
        "XML bomb/vector attacks. Please install defusedxml for security.",
        UserWarning
    )

# DOC íŒŒì¼ ìƒì„±ì„ ìœ„í•œ python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Documentê°€ ì—†ì„ ë•Œë¥¼ ìœ„í•œ ë”ë¯¸ í´ë˜ìŠ¤ (íƒ€ì… íŒíŠ¸ìš©)
    Document = None  # type: ignore
    import warnings
    warnings.warn(
        "python-docx is not installed. DOC report generation will not be available.",
        UserWarning
    )


class ReportGenerator:
    """ë¦¬í¬íŠ¸ ìƒì„±ê¸°"""
    
    def __init__(self):
        pass
    
    def generate_pdf_report(
        self,
        analysis_result: Dict[str, Any],
        product_data: Optional[Dict[str, Any]] = None,
        shop_data: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        PDF ë¦¬í¬íŠ¸ ìƒì„±
        
        Args:
            analysis_result: ë¶„ì„ ê²°ê³¼
            product_data: ìƒí’ˆ ë°ì´í„° (ì„ íƒì‚¬í•­)
            shop_data: Shop ë°ì´í„° (ì„ íƒì‚¬í•­)
            
        Returns:
            PDF íŒŒì¼ ë°”ì´íŠ¸
        """
        # Markdown ë¦¬í¬íŠ¸ë¥¼ ìƒì„±í•˜ì—¬ PDFë¡œ ë³€í™˜
        markdown_content = self.generate_markdown_report(
            analysis_result,
            product_data,
            shop_data
        )
        # ì‹¤ì œ PDF ìƒì„±ì€ reportlab ì‚¬ìš© (í–¥í›„ êµ¬í˜„)
        return markdown_content.encode('utf-8')
    
    def generate_excel_report(
        self,
        analysis_result: Dict[str, Any],
        product_data: Optional[Dict[str, Any]] = None,
        shop_data: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Excel ë¦¬í¬íŠ¸ ìƒì„±
        
        Args:
            analysis_result: ë¶„ì„ ê²°ê³¼
            product_data: ìƒí’ˆ ë°ì´í„° (ì„ íƒì‚¬í•­)
            shop_data: Shop ë°ì´í„° (ì„ íƒì‚¬í•­)
            
        Returns:
            Excel íŒŒì¼ ë°”ì´íŠ¸
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            ws = wb.active
            ws.title = "ë¶„ì„ ë¦¬í¬íŠ¸"
            
            # í—¤ë” ìŠ¤íƒ€ì¼
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            
            row = 1
            
            # ì œëª©
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = "Qoo10 Sales Intelligence Agent - ë¶„ì„ ë¦¬í¬íŠ¸"
            ws[f'A{row}'].font = Font(bold=True, size=16)
            ws[f'A{row}'].alignment = Alignment(horizontal="center", vertical="center")
            row += 2
            
            # ìƒì„±ì¼ì‹œ
            ws[f'A{row}'] = f"ìƒì„±ì¼ì‹œ: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            row += 2
            
            # ìƒí’ˆ ì •ë³´
            if product_data:
                ws[f'A{row}'] = "ìƒí’ˆ ì •ë³´"
                ws[f'A{row}'].font = header_font
                ws[f'A{row}'].fill = header_fill
                row += 1
                
                # ìƒí’ˆëª… (ê°œì„ ëœ ì¶”ì¶œ ë¡œì§ ë°˜ì˜)
                product_name = product_data.get('product_name', 'N/A')
                if product_name and product_name != 'ìƒí’ˆëª… ì—†ìŒ' and product_name != 'N/A':
                    product_name_display = product_name
                else:
                    product_name_display = 'N/A (ì¶”ì¶œ ì‹¤íŒ¨)'
                
                ws[f'A{row}'] = "ìƒí’ˆëª…"
                ws[f'B{row}'] = product_name_display
                row += 1
                ws[f'A{row}'] = "ìƒí’ˆ ì½”ë“œ"
                ws[f'B{row}'] = product_data.get('product_code', 'N/A')
                row += 1
                ws[f'A{row}'] = "ì¹´í…Œê³ ë¦¬"
                ws[f'B{row}'] = product_data.get('category', 'N/A')
                row += 1
                ws[f'A{row}'] = "ë¸Œëœë“œ"
                ws[f'B{row}'] = product_data.get('brand', 'N/A')
                row += 1
                
                # ê°€ê²© ì •ë³´ (ìœ íš¨ì„± ê²€ì¦ëœ ê°’ë§Œ í‘œì‹œ)
                price_data = product_data.get('price', {})
                sale_price = price_data.get('sale_price')
                original_price = price_data.get('original_price')
                
                ws[f'A{row}'] = "íŒë§¤ê°€"
                if sale_price and 100 <= sale_price <= 1000000:
                    ws[f'B{row}'] = f"{sale_price:,}å††"
                else:
                    ws[f'B{row}'] = "N/A"
                row += 1
                
                if original_price and 100 <= original_price <= 1000000:
                    ws[f'A{row}'] = "ì •ê°€"
                    ws[f'B{row}'] = f"{original_price:,}å††"
                    row += 1
                    if sale_price and original_price > sale_price:
                        discount_rate = int((original_price - sale_price) / original_price * 100)
                        ws[f'A{row}'] = "í• ì¸ìœ¨"
                        ws[f'B{row}'] = f"{discount_rate}%"
                        row += 1
                
                # Qãƒã‚¤ãƒ³ãƒˆ ì •ë³´
                qpoint_info = product_data.get('qpoint_info', {})
                if qpoint_info and any(qpoint_info.values()):
                    qpoint_lines = []
                    if qpoint_info.get('max_points'):
                        qpoint_lines.append(f"ìµœëŒ€ {qpoint_info['max_points']}P")
                    if qpoint_info.get('receive_confirmation_points'):
                        qpoint_lines.append(f"ìˆ˜ë ¹í™•ì¸ {qpoint_info['receive_confirmation_points']}P")
                    if qpoint_info.get('review_points'):
                        qpoint_lines.append(f"ë¦¬ë·°ì‘ì„± {qpoint_info['review_points']}P")
                    if qpoint_lines:
                        ws[f'A{row}'] = "Qãƒã‚¤ãƒ³ãƒˆ"
                        ws[f'B{row}'] = ', '.join(qpoint_lines)
                        row += 1
                
                # ë°˜í’ˆ ì •ë³´
                shipping_info = product_data.get('shipping_info', {})
                return_policy = shipping_info.get('return_policy')
                if return_policy:
                    return_text = "ë¬´ë£Œë°˜í’ˆ ê°€ëŠ¥" if return_policy == "free_return" else "ë°˜í’ˆ ê°€ëŠ¥"
                    ws[f'A{row}'] = "ë°˜í’ˆ ì •ì±…"
                    ws[f'B{row}'] = return_text
                    row += 1
                
                row += 1
            
            # Shop ì •ë³´
            if shop_data:
                ws[f'A{row}'] = "Shop ì •ë³´"
                ws[f'A{row}'].font = header_font
                ws[f'A{row}'].fill = header_fill
                row += 1
                
                ws[f'A{row}'] = "Shop ì´ë¦„"
                ws[f'B{row}'] = shop_data.get('shop_name', 'N/A')
                row += 1
                ws[f'A{row}'] = "Shop ë ˆë²¨"
                ws[f'B{row}'] = shop_data.get('shop_level', 'N/A')
                row += 1
                ws[f'A{row}'] = "íŒ”ë¡œì›Œ ìˆ˜"
                ws[f'B{row}'] = shop_data.get('follower_count', 0)
                row += 1
                ws[f'A{row}'] = "ìƒí’ˆ ìˆ˜"
                ws[f'B{row}'] = shop_data.get('product_count', 0)
                row += 2
            
            # ë¶„ì„ ê²°ê³¼
            if "product_analysis" in analysis_result:
                product_analysis = analysis_result["product_analysis"]
                ws[f'A{row}'] = "ìƒí’ˆ ë¶„ì„ ê²°ê³¼"
                ws[f'A{row}'].font = header_font
                ws[f'A{row}'].fill = header_fill
                row += 1
                
                ws[f'A{row}'] = "ì¢…í•© ì ìˆ˜"
                ws[f'B{row}'] = f"{product_analysis.get('overall_score', 0)}/100"
                row += 1
                
                # ê° ë¶„ì„ í•­ëª©
                for analysis_type in ["image_analysis", "description_analysis", "price_analysis", 
                                     "review_analysis", "seo_analysis", "page_structure_analysis"]:
                    if analysis_type in product_analysis:
                        analysis = product_analysis[analysis_type]
                        ws[f'A{row}'] = analysis_type.replace("_", " ").title()
                        ws[f'B{row}'] = f"{analysis.get('score', 0)}/100"
                        row += 1
                row += 1
            
            # ì»¬ëŸ¼ ë„ˆë¹„ ì¡°ì •
            ws.column_dimensions['A'].width = 25
            ws.column_dimensions['B'].width = 40
            
            buffer = io.BytesIO()
            wb.save(buffer)
            return buffer.getvalue()
        except ImportError:
            # openpyxlì´ ì—†ëŠ” ê²½ìš° JSON ë°˜í™˜
            excel_data = {
                "metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "report_type": "excel"
                },
                "analysis_result": analysis_result,
                "product_data": product_data,
                "shop_data": shop_data
            }
            return json.dumps(excel_data, ensure_ascii=False, indent=2).encode('utf-8')
    
    def generate_markdown_report(
        self,
        analysis_result: Dict[str, Any],
        product_data: Optional[Dict[str, Any]] = None,
        shop_data: Optional[Dict[str, Any]] = None,
        validation_result: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Markdown ë¦¬í¬íŠ¸ ìƒì„±
        
        Args:
            analysis_result: ë¶„ì„ ê²°ê³¼
            product_data: ìƒí’ˆ ë°ì´í„° (ì„ íƒì‚¬í•­)
            shop_data: Shop ë°ì´í„° (ì„ íƒì‚¬í•­)
            validation_result: ê²€ì¦ ê²°ê³¼ (ì„ íƒì‚¬í•­)
            
        Returns:
            Markdown ë¬¸ìì—´
        """
        # #region agent log - H5 ê°€ì„¤ ê²€ì¦
        _log_debug("debug-session", "run1", "H5", "report_generator.py:generate_markdown_report", "ë¦¬í¬íŠ¸ ìƒì„± í•¨ìˆ˜ í˜¸ì¶œ - ì²´í¬ë¦¬ìŠ¤íŠ¸ í™•ì¸", {
            "has_analysis_result": bool(analysis_result),
            "analysis_result_keys": list(analysis_result.keys()) if analysis_result and isinstance(analysis_result, dict) else None,
            "has_checklist_in_result": "checklist" in analysis_result if analysis_result and isinstance(analysis_result, dict) else False,
            "checklist_data": analysis_result.get("checklist") if analysis_result and isinstance(analysis_result, dict) else None,
            "checklist_overall_completion": analysis_result.get("checklist", {}).get("overall_completion") if analysis_result and isinstance(analysis_result, dict) and analysis_result.get("checklist") else None,
            "checklist_count": len(analysis_result.get("checklist", {}).get("checklists", [])) if analysis_result and isinstance(analysis_result, dict) and analysis_result.get("checklist") else 0,
            "has_validation_result": bool(validation_result)
        })
        # #endregion
        return self._generate_report_content(
            analysis_result,
            product_data,
            shop_data,
            format="markdown",
            validation_result=validation_result
        )
    
    def generate_doc_report(
        self,
        analysis_result: Dict[str, Any],
        product_data: Optional[Dict[str, Any]] = None,
        shop_data: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        DOC ë¦¬í¬íŠ¸ ìƒì„±
        
        Args:
            analysis_result: ë¶„ì„ ê²°ê³¼
            product_data: ìƒí’ˆ ë°ì´í„° (ì„ íƒì‚¬í•­)
            shop_data: Shop ë°ì´í„° (ì„ íƒì‚¬í•­)
            
        Returns:
            DOC íŒŒì¼ ë°”ì´íŠ¸
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOC report generation")
        
        doc = Document()
        
        # ì œëª©
        title = doc.add_heading('Qoo10 Sales Intelligence Agent - ë¶„ì„ ë¦¬í¬íŠ¸', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ìƒì„±ì¼ì‹œ
        date_para = doc.add_paragraph(f'ìƒì„±ì¼ì‹œ: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # ë¹ˆ ì¤„
        
        # ìƒí’ˆ ì •ë³´
        if product_data:
            doc.add_heading('ìƒí’ˆ ì •ë³´', level=1)
            
            # ìƒí’ˆëª… (ê°œì„ ëœ ì¶”ì¶œ ë¡œì§ ë°˜ì˜)
            product_name = product_data.get('product_name', 'N/A')
            if product_name and product_name != 'ìƒí’ˆëª… ì—†ìŒ' and product_name != 'N/A':
                product_name_display = product_name
            else:
                product_name_display = 'N/A (ì¶”ì¶œ ì‹¤íŒ¨)'
            
            info_items = [
                ("ìƒí’ˆëª…", product_name_display),
                ("ìƒí’ˆ ì½”ë“œ", product_data.get('product_code', 'N/A')),
                ("ì¹´í…Œê³ ë¦¬", product_data.get('category', 'N/A')),
                ("ë¸Œëœë“œ", product_data.get('brand', 'N/A')),
            ]
            
            # ê°€ê²© ì •ë³´ (ìœ íš¨ì„± ê²€ì¦ëœ ê°’ë§Œ í‘œì‹œ)
            price_data = product_data.get('price', {})
            sale_price = price_data.get('sale_price')
            original_price = price_data.get('original_price')
            
            if sale_price and 100 <= sale_price <= 1000000:
                info_items.append(("íŒë§¤ê°€", f"{sale_price:,}å††"))
            else:
                info_items.append(("íŒë§¤ê°€", "N/A"))
            
            if original_price and 100 <= original_price <= 1000000:
                info_items.append(("ì •ê°€", f"{original_price:,}å††"))
                if sale_price and original_price > sale_price:
                    discount_rate = int((original_price - sale_price) / original_price * 100)
                    info_items.append(("í• ì¸ìœ¨", f"{discount_rate}%"))
            
            # Qãƒã‚¤ãƒ³ãƒˆ ì •ë³´
            qpoint_info = product_data.get('qpoint_info', {})
            if qpoint_info and any(qpoint_info.values()):
                qpoint_lines = []
                if qpoint_info.get('max_points'):
                    qpoint_lines.append(f"ìµœëŒ€ {qpoint_info['max_points']}P")
                if qpoint_info.get('receive_confirmation_points'):
                    qpoint_lines.append(f"ìˆ˜ë ¹í™•ì¸ {qpoint_info['receive_confirmation_points']}P")
                if qpoint_info.get('review_points'):
                    qpoint_lines.append(f"ë¦¬ë·°ì‘ì„± {qpoint_info['review_points']}P")
                if qpoint_lines:
                    info_items.append(("Qãƒã‚¤ãƒ³ãƒˆ", ', '.join(qpoint_lines)))
                else:
                    info_items.append(("Qãƒã‚¤ãƒ³ãƒˆ", "N/A"))
            else:
                info_items.append(("Qãƒã‚¤ãƒ³ãƒˆ", "N/A"))
            
            # ë°˜í’ˆ ì •ë³´
            shipping_info = product_data.get('shipping_info', {})
            return_policy = shipping_info.get('return_policy')
            if return_policy:
                return_text = "ë¬´ë£Œë°˜í’ˆ ê°€ëŠ¥" if return_policy == "free_return" else "ë°˜í’ˆ ê°€ëŠ¥"
                info_items.append(("ë°˜í’ˆ ì •ì±…", return_text))
            else:
                info_items.append(("ë°˜í’ˆ ì •ì±…", "N/A"))
            
            self._add_info_table(doc, info_items)
            doc.add_paragraph()
        
        # Shop ì •ë³´
        if shop_data:
            doc.add_heading('Shop ì •ë³´', level=1)
            self._add_info_table(doc, [
                ("Shop ì´ë¦„", shop_data.get('shop_name', 'N/A')),
                ("Shop ë ˆë²¨", shop_data.get('shop_level', 'N/A')),
                ("íŒ”ë¡œì›Œ ìˆ˜", f"{shop_data.get('follower_count', 0):,}ëª…"),
                ("ìƒí’ˆ ìˆ˜", f"{shop_data.get('product_count', 0)}ê°œ"),
            ])
            doc.add_paragraph()
        
        # ìƒí’ˆ ë¶„ì„ ê²°ê³¼
        if "product_analysis" in analysis_result:
            product_analysis = analysis_result["product_analysis"]
            doc.add_heading('ìƒí’ˆ ë¶„ì„ ê²°ê³¼', level=1)
            
            # ì¢…í•© ì ìˆ˜
            overall_score = product_analysis.get('overall_score', 0)
            score_para = doc.add_paragraph()
            score_para.add_run('ì¢…í•© ì ìˆ˜: ').bold = True
            score_para.add_run(f'{overall_score}/100').bold = True
            score_para.add_run(f' ({self._get_grade(overall_score)})')
            doc.add_paragraph()
            
            # ì´ë¯¸ì§€ ë¶„ì„
            self._add_analysis_section(doc, "ì´ë¯¸ì§€ ë¶„ì„", product_analysis.get("image_analysis", {}))
            
            # ì„¤ëª… ë¶„ì„
            self._add_analysis_section(doc, "ìƒí’ˆ ì„¤ëª… ë¶„ì„", product_analysis.get("description_analysis", {}))
            
            # ê°€ê²© ë¶„ì„
            self._add_price_analysis_section(doc, product_analysis.get("price_analysis", {}))
            
            # ë¦¬ë·° ë¶„ì„
            self._add_review_analysis_section(doc, product_analysis.get("review_analysis", {}))
            
            # SEO ë¶„ì„
            self._add_analysis_section(doc, "SEO ë¶„ì„", product_analysis.get("seo_analysis", {}))
            
            # í˜ì´ì§€ êµ¬ì¡° ë¶„ì„
            self._add_analysis_section(doc, "í˜ì´ì§€ êµ¬ì¡° ë¶„ì„", product_analysis.get("page_structure_analysis", {}))
        
        # Shop ë¶„ì„
        if "shop_analysis" in analysis_result:
            shop_analysis = analysis_result["shop_analysis"]
            doc.add_heading('Shop ë¶„ì„ ê²°ê³¼', level=1)
            
            overall_score = shop_analysis.get('overall_score', 0)
            score_para = doc.add_paragraph()
            score_para.add_run('ì¢…í•© ì ìˆ˜: ').bold = True
            score_para.add_run(f'{overall_score}/100').bold = True
            score_para.add_run(f' ({self._get_grade(overall_score)})')
            doc.add_paragraph()
            
            # Shop ì •ë³´ ë¶„ì„
            self._add_shop_info_section(doc, shop_analysis)
            
            # Shop íŠ¹ìˆ˜ì„± ë¶„ì„
            if "shop_specialty" in shop_analysis:
                self._add_shop_specialty_section(doc, shop_analysis.get("shop_specialty", {}))
            
            # ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸
            if "customized_insights" in shop_analysis:
                self._add_customized_insights_section(doc, shop_analysis.get("customized_insights", {}))
        
        # ì¶”ì²œ ì•„ì´ë””ì–´
        recommendations = analysis_result.get("recommendations", [])
        if recommendations:
            doc.add_heading('ë§¤ì¶œ ê°•í™” ì•„ì´ë””ì–´', level=1)
            for i, rec in enumerate(recommendations, 1):
                priority = rec.get("priority", "medium").upper()
                title_text = f"{i}. [{priority}] {rec.get('title', 'N/A')}"
                doc.add_heading(title_text, level=2)
                
                doc.add_paragraph(rec.get('description', 'N/A'))
                
                if rec.get("action_items"):
                    doc.add_paragraph('ì‹¤í–‰ ë°©ë²•:', style='List Bullet')
                    for item in rec["action_items"]:
                        doc.add_paragraph(item, style='List Bullet 2')
                doc.add_paragraph()
        
        # ì²´í¬ë¦¬ìŠ¤íŠ¸
        checklist = analysis_result.get("checklist", {})
        if checklist:
            doc.add_heading('ë©”ë‰´ì–¼ ê¸°ë°˜ ì²´í¬ë¦¬ìŠ¤íŠ¸', level=1)
            overall_completion = checklist.get('overall_completion', 0)
            doc.add_paragraph(f'ì „ì²´ ì™„ì„±ë„: {overall_completion}%')
            doc.add_paragraph()
            
            for cl in checklist.get("checklists", []):
                doc.add_heading(f"{cl.get('category', 'N/A')}: {cl.get('completion_rate', 0)}%", level=2)
                for item in cl.get("items", []):
                    status = "âœ…" if item.get("status") == "completed" else "â¬œ"
                    doc.add_paragraph(f"{status} {item.get('title', 'N/A')}", style='List Bullet')
                doc.add_paragraph()
        
        # ê²½ìŸì‚¬ ë¶„ì„
        competitor_analysis = analysis_result.get("competitor_analysis", {})
        if competitor_analysis:
            doc.add_heading('ê²½ìŸì‚¬ ë¹„êµ ë¶„ì„', level=1)
            comparison = competitor_analysis.get("comparison", {})
            doc.add_paragraph(f'ê°€ê²© í¬ì§€ì…”ë‹: {comparison.get("price_position", "N/A")}')
            doc.add_paragraph(f'í‰ì  í¬ì§€ì…”ë‹: {comparison.get("rating_position", "N/A")}')
            doc.add_paragraph(f'ë¦¬ë·° í¬ì§€ì…”ë‹: {comparison.get("review_position", "N/A")}')
            doc.add_paragraph()
            
            if competitor_analysis.get("differentiation_points"):
                doc.add_heading('ì°¨ë³„í™” í¬ì¸íŠ¸', level=2)
                for point in competitor_analysis["differentiation_points"]:
                    doc.add_paragraph(point, style='List Bullet')
        
        # ë¬¸ì„œë¥¼ ë°”ì´íŠ¸ë¡œ ë³€í™˜
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def _add_info_table(self, doc: Any, data: List[tuple]):
        """ì •ë³´ í…Œì´ë¸” ì¶”ê°€"""
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (key, value) in enumerate(data):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    def _add_analysis_section(self, doc: Any, title: str, analysis: Dict[str, Any]):
        """ë¶„ì„ ì„¹ì…˜ ì¶”ê°€"""
        doc.add_heading(title, level=2)
        
        score = analysis.get('score', 0)
        score_para = doc.add_paragraph()
        score_para.add_run('ì ìˆ˜: ').bold = True
        score_para.add_run(f'{score}/100').bold = True
        score_para.add_run(f' ({self._get_grade(score)})')
        
        # ì„¸ë¶€ ì •ë³´ ì¶”ê°€
        if title == "ì´ë¯¸ì§€ ë¶„ì„":
            doc.add_paragraph(f'ì¸ë„¤ì¼ í’ˆì§ˆ: {analysis.get("thumbnail_quality", "N/A")}')
            doc.add_paragraph(f'ìƒì„¸ ì´ë¯¸ì§€ ê°œìˆ˜: {analysis.get("image_count", 0)}ê°œ')
        elif title == "ìƒí’ˆ ì„¤ëª… ë¶„ì„":
            doc.add_paragraph(f'ì„¤ëª… ê¸¸ì´: {analysis.get("description_length", 0)}ì')
            doc.add_paragraph(f'êµ¬ì¡°í™” í’ˆì§ˆ: {analysis.get("structure_quality", "N/A")}')
        elif title == "SEO ë¶„ì„":
            doc.add_paragraph(f'í‚¤ì›Œë“œ ìƒí’ˆëª… í¬í•¨: {"ì˜ˆ" if analysis.get("keywords_in_name") else "ì•„ë‹ˆì˜¤"}')
            doc.add_paragraph(f'í‚¤ì›Œë“œ ì„¤ëª… í¬í•¨: {"ì˜ˆ" if analysis.get("keywords_in_description") else "ì•„ë‹ˆì˜¤"}')
            doc.add_paragraph(f'ì¹´í…Œê³ ë¦¬ ì„¤ì •: {"ì˜ˆ" if analysis.get("category_set") else "ì•„ë‹ˆì˜¤"}')
            doc.add_paragraph(f'ë¸Œëœë“œ ì„¤ì •: {"ì˜ˆ" if analysis.get("brand_set") else "ì•„ë‹ˆì˜¤"}')
        elif title == "í˜ì´ì§€ êµ¬ì¡° ë¶„ì„":
            doc.add_paragraph(f'ì „ì²´ í´ë˜ìŠ¤ ìˆ˜: {analysis.get("total_classes", 0)}ê°œ')
            key_elements = analysis.get("key_elements_present", {})
            if key_elements:
                doc.add_paragraph('ì£¼ìš” ìš”ì†Œ ì¡´ì¬ ì—¬ë¶€:')
                for key, present in key_elements.items():
                    doc.add_paragraph(f'  - {key}: {"ì˜ˆ" if present else "ì•„ë‹ˆì˜¤"}', style='List Bullet 2')
        
        # ì¶”ì²œ ì‚¬í•­
        if analysis.get("recommendations"):
            doc.add_paragraph('ì¶”ì²œ ì‚¬í•­:', style='List Bullet')
            for rec in analysis["recommendations"]:
                doc.add_paragraph(rec, style='List Bullet 2')
        
        doc.add_paragraph()
    
    def _add_price_analysis_section(self, doc: Any, analysis: Dict[str, Any]):
        """ê°€ê²© ë¶„ì„ ì„¹ì…˜ ì¶”ê°€ (ê°œì„ ëœ í¬ë¡¤ëŸ¬ ë°ì´í„° ë°˜ì˜)"""
        doc.add_heading('ê°€ê²© ë¶„ì„', level=2)
        
        score = analysis.get('score', 0)
        score_para = doc.add_paragraph()
        score_para.add_run('ì ìˆ˜: ').bold = True
        score_para.add_run(f'{score}/100').bold = True
        score_para.add_run(f' ({self._get_grade(score)})')
        
        # ìœ íš¨ì„± ê²€ì¦ëœ ê°€ê²©ë§Œ í‘œì‹œ (100~1,000,000ì—” ë²”ìœ„)
        sale_price = analysis.get('sale_price')
        original_price = analysis.get('original_price')
        discount_rate = analysis.get('discount_rate', 0) or 0
        
        if sale_price and 100 <= sale_price <= 1000000:
            doc.add_paragraph(f'íŒë§¤ê°€: {sale_price:,}å††')
        else:
            doc.add_paragraph('íŒë§¤ê°€: N/A (ìœ íš¨í•˜ì§€ ì•Šì€ ê°’)')
        
        if original_price and 100 <= original_price <= 1000000:
            doc.add_paragraph(f'ì •ê°€: {original_price:,}å††')
            if sale_price and original_price > sale_price:
                calculated_discount = int((original_price - sale_price) / original_price * 100)
                doc.add_paragraph(f'í• ì¸ìœ¨: {calculated_discount}%')
        elif discount_rate > 0:
            doc.add_paragraph(f'í• ì¸ìœ¨: {discount_rate}%')
        
        positioning = analysis.get('positioning', '')
        if positioning:
            doc.add_paragraph(f'ê°€ê²© í¬ì§€ì…”ë‹: {positioning}')
        
        if analysis.get("recommendations"):
            doc.add_paragraph('ì¶”ì²œ ì‚¬í•­:', style='List Bullet')
            for rec in analysis["recommendations"]:
                doc.add_paragraph(rec, style='List Bullet 2')
        
        doc.add_paragraph()
    
    def _add_review_analysis_section(self, doc: Any, analysis: Dict[str, Any]):
        """ë¦¬ë·° ë¶„ì„ ì„¹ì…˜ ì¶”ê°€ (ê°œì„ ëœ í¬ë¡¤ëŸ¬ ë°ì´í„° ë°˜ì˜)"""
        doc.add_heading('ë¦¬ë·° ë¶„ì„', level=2)
        
        score = analysis.get('score', 0)
        score_para = doc.add_paragraph()
        score_para.add_run('ì ìˆ˜: ').bold = True
        score_para.add_run(f'{score}/100').bold = True
        score_para.add_run(f' ({self._get_grade(score)})')
        
        rating = analysis.get('rating', 0) or 0.0
        review_count = analysis.get('review_count', 0) or 0
        # fallback: reviews ë°°ì—´ ê¸¸ì´ ì‚¬ìš©
        reviews_list = analysis.get('reviews', [])
        if review_count == 0 and len(reviews_list) > 0:
            review_count = len(reviews_list)
        
        negative_ratio = analysis.get('negative_ratio', 0.0) or 0.0
        
        doc.add_paragraph(f'í‰ì : {rating:.1f}/5.0')
        if review_count > 0:
            doc.add_paragraph(f'ë¦¬ë·° ìˆ˜: {review_count:,}ê°œ')
        else:
            doc.add_paragraph('ë¦¬ë·° ìˆ˜: 0ê°œ (ë˜ëŠ” ì¶”ì¶œ ì‹¤íŒ¨)')
        
        if len(reviews_list) > 0:
            doc.add_paragraph(f'ì¶”ì¶œëœ ë¦¬ë·° í…ìŠ¤íŠ¸: {len(reviews_list)}ê°œ')
        
        if negative_ratio > 0:
            doc.add_paragraph(f'ë¶€ì • ë¦¬ë·° ë¹„ìœ¨: {negative_ratio:.1%}')
        
        if analysis.get("recommendations"):
            doc.add_paragraph('ì¶”ì²œ ì‚¬í•­:', style='List Bullet')
            for rec in analysis["recommendations"]:
                doc.add_paragraph(rec, style='List Bullet 2')
        
        doc.add_paragraph()
    
    def _add_shop_info_section(self, doc: Any, shop_analysis: Dict[str, Any]):
        """Shop ì •ë³´ ë¶„ì„ ì„¹ì…˜ ì¶”ê°€"""
        shop_info = shop_analysis.get("shop_info", {})
        if shop_info:
            doc.add_heading('Shop ì •ë³´ ë¶„ì„', level=2)
            doc.add_paragraph(f'ì ìˆ˜: {shop_info.get("score", 0)}/100 ({self._get_grade(shop_info.get("score", 0))})')
            doc.add_paragraph()
        
        level_analysis = shop_analysis.get("level_analysis", {})
        if level_analysis:
            doc.add_heading('Shop ë ˆë²¨ ë¶„ì„', level=2)
            doc.add_paragraph(f'í˜„ì¬ ë ˆë²¨: {level_analysis.get("current_level", "N/A")}')
            doc.add_paragraph(f'ì •ì‚° ë¦¬ë“œíƒ€ì„: {level_analysis.get("settlement_leadtime", 15)}ì¼')
            doc.add_paragraph(f'ëª©í‘œ ë ˆë²¨: {level_analysis.get("target_level", "N/A")}')
            
            if level_analysis.get("requirements"):
                doc.add_paragraph('ìš”êµ¬ì‚¬í•­:', style='List Bullet')
                for req in level_analysis["requirements"]:
                    doc.add_paragraph(req, style='List Bullet 2')
            
            if level_analysis.get("recommendations"):
                doc.add_paragraph('ì¶”ì²œ ì‚¬í•­:', style='List Bullet')
                for rec in level_analysis["recommendations"]:
                    doc.add_paragraph(rec, style='List Bullet 2')
            doc.add_paragraph()
    
    def _add_shop_specialty_section(self, doc: Any, specialty: Dict[str, Any]):
        """Shop íŠ¹ìˆ˜ì„± ì„¹ì…˜ ì¶”ê°€"""
        doc.add_heading('Shop íŠ¹ìˆ˜ì„± ë¶„ì„', level=2)
        
        doc.add_paragraph(f'ë¸Œëœë“œ ìƒµ ì—¬ë¶€: {"ì˜ˆ" if specialty.get("is_brand_shop") else "ì•„ë‹ˆì˜¤"}')
        if specialty.get("brand_name"):
            doc.add_paragraph(f'ë¸Œëœë“œëª…: {specialty.get("brand_name")}')
        
        lineup_type = specialty.get("product_lineup_type", "mixed")
        doc.add_paragraph(f'ì œí’ˆ ë¼ì¸ì—… íŠ¹ì„±: {lineup_type}')
        
        target_customer = specialty.get("target_customer", "general")
        doc.add_paragraph(f'íƒ€ê²Ÿ ê³ ê°ì¸µ: {target_customer.replace("_", " ")}')
        
        unique_features = specialty.get("unique_features", [])
        if unique_features:
            doc.add_paragraph('ë…íŠ¹í•œ íŠ¹ì§•:', style='List Bullet')
            for feature in unique_features:
                doc.add_paragraph(feature, style='List Bullet 2')
        
        doc.add_paragraph(f'íŠ¹ìˆ˜ì„± ì ìˆ˜: {specialty.get("specialty_score", 0)}/100')
        doc.add_paragraph()
    
    def _add_customized_insights_section(self, doc: Any, insights: Dict[str, Any]):
        """ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸ ì„¹ì…˜ ì¶”ê°€"""
        doc.add_heading('ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸', level=2)
        
        if insights.get("shop_positioning"):
            doc.add_paragraph(f'Shop í¬ì§€ì…”ë‹: {insights.get("shop_positioning")}')
            doc.add_paragraph()
        
        strengths = insights.get("strengths", [])
        if strengths:
            doc.add_heading('ê°•ì ', level=3)
            for strength in strengths:
                doc.add_paragraph(strength, style='List Bullet')
            doc.add_paragraph()
        
        opportunities = insights.get("opportunities", [])
        if opportunities:
            doc.add_heading('ê¸°íšŒ', level=3)
            for opp in opportunities:
                doc.add_paragraph(opp, style='List Bullet')
            doc.add_paragraph()
        
        recommendations = insights.get("recommendations", [])
        if recommendations:
            doc.add_heading('ì¶”ì²œ ì‚¬í•­', level=3)
            for rec in recommendations:
                doc.add_paragraph(rec, style='List Bullet')
            doc.add_paragraph()
        
        advantages = insights.get("competitive_advantages", [])
        if advantages:
            doc.add_heading('ê²½ìŸ ìš°ìœ„', level=3)
            for adv in advantages:
                doc.add_paragraph(adv, style='List Bullet')
            doc.add_paragraph()
    
    def _get_grade(self, score: int) -> str:
        """ì ìˆ˜ì— ë”°ë¥¸ ë“±ê¸‰ ë°˜í™˜"""
        if score >= 90:
            return "Excellent"
        elif score >= 70:
            return "Good"
        elif score >= 50:
            return "Fair"
        else:
            return "Poor"
    
    def _generate_report_content(
        self,
        analysis_result: Dict[str, Any],
        product_data: Optional[Dict[str, Any]],
        shop_data: Optional[Dict[str, Any]],
        format: str = "markdown",
        validation_result: Optional[Dict[str, Any]] = None
    ) -> str:
        # #region agent log - H3 ê°€ì„¤ ê²€ì¦
        _log_debug("debug-session", "run1", "H3", "report_generator.py:_generate_report_content", "ë¦¬í¬íŠ¸ ìƒì„± ì‹œì‘ - ì…ë ¥ ë°ì´í„° êµ¬ì¡°", {
            "has_analysis_result": bool(analysis_result),
            "analysis_result_keys": list(analysis_result.keys()) if analysis_result and isinstance(analysis_result, dict) else None,
            "has_product_analysis": "product_analysis" in analysis_result if analysis_result and isinstance(analysis_result, dict) else False,
            "has_checklist": "checklist" in analysis_result if analysis_result and isinstance(analysis_result, dict) else False,
            "has_product_data": bool(product_data),
            "product_data_keys": list(product_data.keys()) if product_data and isinstance(product_data, dict) else None,
            "product_name_in_data": product_data.get("product_name") if product_data and isinstance(product_data, dict) else None,
            "price_sale_in_data": product_data.get("price", {}).get("sale_price") if product_data and isinstance(product_data, dict) and product_data.get("price") else None
        })
        # #endregion
        """ë¦¬í¬íŠ¸ ë‚´ìš© ìƒì„± (Markdown í˜•ì‹)"""
        lines = []
        
        # í—¤ë”
        lines.append("# Qoo10 Sales Intelligence Agent - ë¶„ì„ ë¦¬í¬íŠ¸")
        lines.append(f"\n**ìƒì„±ì¼ì‹œ:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("\n---\n")
        
        # ë°ì´í„° ì†ŒìŠ¤ í‘œì‹œ (í¬ë¡¤ë§ ë°©ë²• ë˜ëŠ” API)
        crawled_with = None
        if product_data and "crawled_with" in product_data:
            crawled_with = product_data["crawled_with"]
        elif shop_data and "crawled_with" in shop_data:
            crawled_with = shop_data["crawled_with"]
        
        if crawled_with:
            if crawled_with == "qoo10_api":
                lines.append(f"**ë°ì´í„° ì†ŒìŠ¤:** Qoo10 ê³µì‹ API")
            else:
                lines.append(f"**í¬ë¡¤ë§ ë°©ë²•:** {crawled_with.upper()}")
            lines.append("\n---\n")
        
        # ìƒí’ˆ ì •ë³´
        if product_data:
            # #region agent log - H3 ê°€ì„¤ ê²€ì¦
            _log_debug("debug-session", "run1", "H3", "report_generator.py:_generate_report_content", "ìƒí’ˆ ì •ë³´ ë¦¬í¬íŠ¸ì— ì¶”ê°€ ì‹œì‘", {
                "product_name": product_data.get('product_name'),
                "price_sale": product_data.get('price', {}).get('sale_price'),
                "price_original": product_data.get('price', {}).get('original_price'),
                "has_qpoint": bool(product_data.get('qpoint_info')),
                "has_coupon": bool(product_data.get('coupon_info', {}).get('has_coupon'))
            })
            # #endregion
            lines.append("## ğŸ“¦ ìƒí’ˆ ì •ë³´")
            lines.append("")
            lines.append("| í•­ëª© | ë‚´ìš© |")
            lines.append("|------|------|")
            
            # ìƒí’ˆëª… (ê°œì„ ëœ ì¶”ì¶œ ë¡œì§ ë°˜ì˜)
            product_name = product_data.get('product_name', 'N/A')
            if product_name and product_name != 'ìƒí’ˆëª… ì—†ìŒ' and product_name != 'N/A':
                lines.append(f"| ìƒí’ˆëª… | {product_name} |")
            else:
                lines.append(f"| ìƒí’ˆëª… | N/A (ì¶”ì¶œ ì‹¤íŒ¨) |")
            
            lines.append(f"| ìƒí’ˆ ì½”ë“œ | {product_data.get('product_code', 'N/A')} |")
            lines.append(f"| ì¹´í…Œê³ ë¦¬ | {product_data.get('category', 'N/A')} |")
            lines.append(f"| ë¸Œëœë“œ | {product_data.get('brand', 'N/A')} |")
            
            # ê°€ê²© ì •ë³´ (ìœ íš¨ì„± ê²€ì¦ëœ ê°’ë§Œ í‘œì‹œ)
            price_data = product_data.get('price', {})
            sale_price = price_data.get('sale_price')
            original_price = price_data.get('original_price')
            
            if sale_price and 100 <= sale_price <= 1000000:  # ìœ íš¨ì„± ê²€ì¦
                lines.append(f"| íŒë§¤ê°€ | {sale_price:,}å†† |")
            else:
                lines.append(f"| íŒë§¤ê°€ | N/A |")
            
            if original_price and 100 <= original_price <= 1000000:  # ìœ íš¨ì„± ê²€ì¦
                lines.append(f"| ì •ê°€ | {original_price:,}å†† |")
                if sale_price and original_price > sale_price:
                    discount_rate = int((original_price - sale_price) / original_price * 100)
                    lines.append(f"| í• ì¸ìœ¨ | {discount_rate}% |")
            
            # Qãƒã‚¤ãƒ³ãƒˆ ì •ë³´ (ê°œì„ ëœ ì¶”ì¶œ ë¡œì§ ë°˜ì˜)
            qpoint_info = product_data.get('qpoint_info', {})
            if qpoint_info and any(qpoint_info.values()):
                qpoint_lines = []
                if qpoint_info.get('max_points'):
                    qpoint_lines.append(f"ìµœëŒ€ {qpoint_info['max_points']}P")
                if qpoint_info.get('receive_confirmation_points'):
                    qpoint_lines.append(f"ìˆ˜ë ¹í™•ì¸ {qpoint_info['receive_confirmation_points']}P")
                if qpoint_info.get('review_points'):
                    qpoint_lines.append(f"ë¦¬ë·°ì‘ì„± {qpoint_info['review_points']}P")
                if qpoint_info.get('auto_points'):
                    qpoint_lines.append(f"ìë™ {qpoint_info['auto_points']}P")
                
                if qpoint_lines:
                    lines.append(f"| Qãƒã‚¤ãƒ³ãƒˆ | {', '.join(qpoint_lines)} |")
                else:
                    lines.append(f"| Qãƒã‚¤ãƒ³ãƒˆ | N/A |")
            else:
                lines.append(f"| Qãƒã‚¤ãƒ³ãƒˆ | N/A |")
            
            # ë°˜í’ˆ ì •ë³´ (ê°œì„ ëœ ì¶”ì¶œ ë¡œì§ ë°˜ì˜)
            shipping_info = product_data.get('shipping_info', {})
            return_policy = shipping_info.get('return_policy')
            if return_policy:
                return_text = "ë¬´ë£Œë°˜í’ˆ ê°€ëŠ¥" if return_policy == "free_return" else "ë°˜í’ˆ ê°€ëŠ¥"
                lines.append(f"| ë°˜í’ˆ ì •ì±… | {return_text} |")
            else:
                lines.append(f"| ë°˜í’ˆ ì •ì±… | N/A |")
            
            # ë°°ì†¡ ì •ë³´
            if shipping_info.get('free_shipping'):
                lines.append(f"| ë°°ì†¡ | ë¬´ë£Œë°°ì†¡ |")
            elif shipping_info.get('shipping_fee'):
                lines.append(f"| ë°°ì†¡ë¹„ | {shipping_info['shipping_fee']:,}å†† |")
            
            # ì¿ í° ì •ë³´
            coupon_info = product_data.get('coupon_info', {})
            if coupon_info.get('has_coupon'):
                coupon_type = coupon_info.get('coupon_type', 'auto')
                max_discount = coupon_info.get('max_discount')
                if max_discount:
                    lines.append(f"| ì¿ í° | {coupon_type} (ìµœëŒ€ {max_discount}å†† í• ì¸) |")
                else:
                    lines.append(f"| ì¿ í° | {coupon_type} |")
            
            lines.append("\n")
        
        # Shop ì •ë³´
        if shop_data:
            lines.append("## ğŸª Shop ì •ë³´")
            lines.append("")
            lines.append("| í•­ëª© | ë‚´ìš© |")
            lines.append("|------|------|")
            lines.append(f"| Shop ì´ë¦„ | {shop_data.get('shop_name', 'N/A')} |")
            lines.append(f"| Shop ë ˆë²¨ | {shop_data.get('shop_level', 'N/A')} |")
            lines.append(f"| íŒ”ë¡œì›Œ ìˆ˜ | {shop_data.get('follower_count', 0):,}ëª… |")
            lines.append(f"| ìƒí’ˆ ìˆ˜ | {shop_data.get('product_count', 0)}ê°œ |")
            lines.append("\n")
        
        # ìƒí’ˆ ë¶„ì„ ê²°ê³¼
        if "product_analysis" in analysis_result:
            product_analysis = analysis_result["product_analysis"]
            lines.append("## ğŸ“Š ìƒí’ˆ ë¶„ì„ ê²°ê³¼")
            lines.append("")
            
            overall_score = product_analysis.get('overall_score', 0)
            grade = self._get_grade(overall_score)
            lines.append(f"### ì¢…í•© ì ìˆ˜: **{overall_score}/100** ({grade})")
            lines.append("")
            
            # ì´ë¯¸ì§€ ë¶„ì„
            self._add_markdown_analysis_section(lines, "ì´ë¯¸ì§€ ë¶„ì„", product_analysis.get("image_analysis", {}))
            
            # ì„¤ëª… ë¶„ì„
            self._add_markdown_analysis_section(lines, "ìƒí’ˆ ì„¤ëª… ë¶„ì„", product_analysis.get("description_analysis", {}))
            
            # ê°€ê²© ë¶„ì„
            self._add_markdown_price_analysis(lines, product_analysis.get("price_analysis", {}))
            
            # ë¦¬ë·° ë¶„ì„
            self._add_markdown_review_analysis(lines, product_analysis.get("review_analysis", {}))
            
            # SEO ë¶„ì„
            self._add_markdown_analysis_section(lines, "SEO ë¶„ì„", product_analysis.get("seo_analysis", {}))
            
            # í˜ì´ì§€ êµ¬ì¡° ë¶„ì„
            self._add_markdown_analysis_section(lines, "í˜ì´ì§€ êµ¬ì¡° ë¶„ì„", product_analysis.get("page_structure_analysis", {}))
        
        # Shop ë¶„ì„
        if "shop_analysis" in analysis_result:
            shop_analysis = analysis_result["shop_analysis"]
            lines.append("## ğŸ¬ Shop ë¶„ì„ ê²°ê³¼")
            lines.append("")
            
            overall_score = shop_analysis.get('overall_score', 0)
            grade = self._get_grade(overall_score)
            lines.append(f"### ì¢…í•© ì ìˆ˜: **{overall_score}/100** ({grade})")
            lines.append("")
            
            # Shop ì •ë³´ ë¶„ì„
            self._add_markdown_shop_info(lines, shop_analysis)
            
            # Shop íŠ¹ìˆ˜ì„± ë¶„ì„
            if "shop_specialty" in shop_analysis:
                self._add_markdown_shop_specialty(lines, shop_analysis.get("shop_specialty", {}))
            
            # ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸
            if "customized_insights" in shop_analysis:
                self._add_markdown_customized_insights(lines, shop_analysis.get("customized_insights", {}))
        
        # AI ì¸ì‚¬ì´íŠ¸ (Gemini ìƒì„±)
        product_analysis = analysis_result.get("product_analysis", {})
        ai_insights = product_analysis.get("ai_insights")
        if ai_insights:
            lines.append("## ğŸ¤– AI ì¸ì‚¬ì´íŠ¸ (Gemini)")
            lines.append("")
            
            strengths = ai_insights.get("strengths", [])
            if strengths:
                lines.append("### ê°•ì ")
                for strength in strengths:
                    lines.append(f"- âœ… {strength}")
                lines.append("")
            
            weaknesses = ai_insights.get("weaknesses", [])
            if weaknesses:
                lines.append("### ê°œì„  í•„ìš” ì‚¬í•­")
                for weakness in weaknesses:
                    lines.append(f"- âš ï¸ {weakness}")
                lines.append("")
            
            action_items = ai_insights.get("action_items", [])
            if action_items:
                lines.append("### ìš°ì„ ìˆœìœ„ ì•¡ì…˜ ì•„ì´í…œ")
                for i, item in enumerate(action_items[:5], 1):  # ìƒìœ„ 5ê°œë§Œ
                    priority = item.get("priority", "medium").upper()
                    priority_emoji = {"HIGH": "ğŸ”´", "MEDIUM": "ğŸŸ¡", "LOW": "ğŸŸ¢"}.get(priority, "âšª")
                    lines.append(f"{i}. {priority_emoji} **{item.get('title', 'N/A')}**")
                    lines.append(f"   - {item.get('description', 'N/A')}")
                    if item.get("expected_impact"):
                        lines.append(f"   - ì˜ˆìƒ íš¨ê³¼: {item.get('expected_impact')}")
                    lines.append("")
            
            insights = ai_insights.get("insights")
            if insights:
                lines.append("### ì¢…í•© ì¸ì‚¬ì´íŠ¸")
                lines.append(insights)
                lines.append("")
        
        # ì¶”ì²œ ì•„ì´ë””ì–´
        recommendations = analysis_result.get("recommendations", [])
        if recommendations:
            lines.append("## ğŸ’¡ ë§¤ì¶œ ê°•í™” ì•„ì´ë””ì–´")
            lines.append("")
            for i, rec in enumerate(recommendations, 1):
                priority = rec.get("priority", "medium").upper()
                priority_emoji = {"HIGH": "ğŸ”´", "MEDIUM": "ğŸŸ¡", "LOW": "ğŸŸ¢"}.get(priority, "âšª")
                lines.append(f"### {i}. {priority_emoji} [{priority}] {rec.get('title', 'N/A')}")
                lines.append("")
                lines.append(rec.get('description', 'N/A'))
                lines.append("")
                if rec.get("action_items"):
                    lines.append("**ì‹¤í–‰ ë°©ë²•:**")
                    for item in rec["action_items"]:
                        lines.append(f"- {item}")
                lines.append("")
        
        # ì²´í¬ë¦¬ìŠ¤íŠ¸
        checklist = analysis_result.get("checklist", {})
        # #region agent log - H5 ê°€ì„¤ ê²€ì¦
        _log_debug("debug-session", "run1", "H5", "report_generator.py:_generate_report_content", "ì²´í¬ë¦¬ìŠ¤íŠ¸ ë°ì´í„° í™•ì¸", {
            "has_checklist": bool(checklist),
            "checklist_keys": list(checklist.keys()) if checklist and isinstance(checklist, dict) else None,
            "overall_completion": checklist.get('overall_completion') if checklist and isinstance(checklist, dict) else None,
            "checklist_count": len(checklist.get("checklists", [])) if checklist and isinstance(checklist, dict) else 0,
            "first_checklist_category": checklist.get("checklists", [{}])[0].get("category") if checklist and isinstance(checklist, dict) and checklist.get("checklists") else None,
            "first_checklist_items_count": len(checklist.get("checklists", [{}])[0].get("items", [])) if checklist and isinstance(checklist, dict) and checklist.get("checklists") else 0,
            "total_items": sum(len(cl.get("items", [])) for cl in checklist.get("checklists", [])) if checklist and isinstance(checklist, dict) else 0,
            "completed_items": sum(len([item for item in cl.get("items", []) if item.get("status") == "completed"]) for cl in checklist.get("checklists", [])) if checklist and isinstance(checklist, dict) else 0
        })
        # #endregion
        if checklist:
            lines.append("## âœ… ë©”ë‰´ì–¼ ê¸°ë°˜ ì²´í¬ë¦¬ìŠ¤íŠ¸")
            lines.append("")
            overall_completion = checklist.get('overall_completion', 0)
            lines.append(f"### ì „ì²´ ì™„ì„±ë„: **{overall_completion}%**")
            lines.append("")
            # #region agent log - H5 ê°€ì„¤ ê²€ì¦
            _log_debug("debug-session", "run1", "H5", "report_generator.py:_generate_report_content", "ì²´í¬ë¦¬ìŠ¤íŠ¸ ë¦¬í¬íŠ¸ì— ì¶”ê°€ ì‹œì‘", {
                "overall_completion": overall_completion,
                "checklist_categories_count": len(checklist.get("checklists", []))
            })
            # #endregion
            items_added_count = 0
            for cl in checklist.get("checklists", []):
                category = cl.get('category', 'N/A')
                completion_rate = cl.get('completion_rate', 0)
                lines.append(f"#### {category}: {completion_rate}%")
                lines.append("")
                for item in cl.get("items", []):
                    status = "âœ…" if item.get("status") == "completed" else "â¬œ"
                    item_title = item.get('title', 'N/A')
                    lines.append(f"- {status} {item_title}")
                    items_added_count += 1
                lines.append("")
            # #region agent log - H5 ê°€ì„¤ ê²€ì¦
            _log_debug("debug-session", "run1", "H5", "report_generator.py:_generate_report_content", "ì²´í¬ë¦¬ìŠ¤íŠ¸ ë¦¬í¬íŠ¸ì— ì¶”ê°€ ì™„ë£Œ", {
                "items_added_to_report": items_added_count,
                "total_items_in_checklist": sum(len(cl.get("items", [])) for cl in checklist.get("checklists", []))
            })
            # #endregion
        else:
            # #region agent log - H5 ê°€ì„¤ ê²€ì¦
            _log_debug("debug-session", "run1", "H5", "report_generator.py:_generate_report_content", "ì²´í¬ë¦¬ìŠ¤íŠ¸ ì—†ìŒ - ë¦¬í¬íŠ¸ì— ì¶”ê°€ë˜ì§€ ì•ŠìŒ", {
                "checklist_in_result": bool(analysis_result.get("checklist"))
            })
            # #endregion
        
        # ê²½ìŸì‚¬ ë¶„ì„
        competitor_analysis = analysis_result.get("competitor_analysis", {})
        if competitor_analysis:
            lines.append("## ğŸ† ê²½ìŸì‚¬ ë¹„êµ ë¶„ì„")
            lines.append("")
            comparison = competitor_analysis.get("comparison", {})
            lines.append(f"### ê°€ê²© í¬ì§€ì…”ë‹: {comparison.get('price_position', 'N/A')}")
            lines.append(f"### í‰ì  í¬ì§€ì…”ë‹: {comparison.get('rating_position', 'N/A')}")
            lines.append(f"### ë¦¬ë·° í¬ì§€ì…”ë‹: {comparison.get('review_position', 'N/A')}")
            lines.append("")
            if competitor_analysis.get("differentiation_points"):
                lines.append("### ì°¨ë³„í™” í¬ì¸íŠ¸:")
                for point in competitor_analysis["differentiation_points"]:
                    lines.append(f"- {point}")
                lines.append("")
        
        # ë°ì´í„° ê²€ì¦ ê²°ê³¼
        if validation_result:
            lines.append("## ğŸ” ë°ì´í„° ê²€ì¦ ê²°ê³¼")
            lines.append("")
            
            validation_score = validation_result.get("validation_score", 0)
            is_valid = validation_result.get("is_valid", False)
            mismatches = validation_result.get("mismatches", [])
            missing_items = validation_result.get("missing_items", [])
            corrected_fields = validation_result.get("corrected_fields", [])
            
            # ê²€ì¦ ì ìˆ˜ ë° ìƒíƒœ
            status_emoji = "âœ…" if is_valid else "âš ï¸"
            status_text = "ì¼ì¹˜" if is_valid else "ë¶ˆì¼ì¹˜"
            lines.append(f"### {status_emoji} ê²€ì¦ ì ìˆ˜: **{validation_score:.1f}%** ({status_text})")
            lines.append("")
            
            # ë³´ì •ëœ í•„ë“œ
            if corrected_fields:
                lines.append(f"**ìë™ ë³´ì •ëœ í•„ë“œ ({len(corrected_fields)}ê°œ):**")
                for field in corrected_fields:
                    lines.append(f"- {field}")
                lines.append("")
            
            # ë¶ˆì¼ì¹˜ í•­ëª©
            if mismatches:
                lines.append(f"**ë¶ˆì¼ì¹˜ í•­ëª© ({len(mismatches)}ê°œ):**")
                for mismatch in mismatches:
                    field = mismatch.get("field", "N/A")
                    crawler_value = mismatch.get("crawler_value", "N/A")
                    report_value = mismatch.get("report_value", "N/A")
                    severity = mismatch.get("severity", "medium")
                    corrected = mismatch.get("corrected", False)
                    severity_emoji = {"high": "ğŸ”´", "medium": "ğŸŸ¡", "low": "ğŸŸ¢"}.get(severity, "âšª")
                    corrected_text = " (ìë™ ë³´ì •ë¨)" if corrected else ""
                    lines.append(f"- {severity_emoji} **{field}**: í¬ë¡¤ëŸ¬={crawler_value}, ë¦¬í¬íŠ¸={report_value}{corrected_text}")
                lines.append("")
            
            # ëˆ„ë½ í•­ëª©
            if missing_items:
                lines.append(f"**ëˆ„ë½ í•­ëª© ({len(missing_items)}ê°œ):**")
                for missing in missing_items:
                    field = missing.get("field", "N/A")
                    checklist_item_id = missing.get("checklist_item_id", "N/A")
                    severity = missing.get("severity", "medium")
                    severity_emoji = {"high": "ğŸ”´", "medium": "ğŸŸ¡", "low": "ğŸŸ¢"}.get(severity, "âšª")
                    lines.append(f"- {severity_emoji} **{field}**: ì²´í¬ë¦¬ìŠ¤íŠ¸ í•­ëª©={checklist_item_id}")
                lines.append("")
            
            # ë°ì´í„° ì†ŒìŠ¤ ì •ë³´
            data_source = validation_result.get("data_source", "unknown")
            has_api_data = validation_result.get("has_api_data", False)
            if has_api_data:
                lines.append(f"**ë°ì´í„° ì†ŒìŠ¤:** Qoo10 ê³µì‹ API (ìš°ì„  ì‚¬ìš©)")
            else:
                lines.append(f"**ë°ì´í„° ì†ŒìŠ¤:** {data_source}")
            
            # êµ¬ì¡° ë¹„êµ ê²°ê³¼ (API êµ¬ì¡° ê¸°ë°˜)
            structure_comparison = validation_result.get("structure_comparison")
            if structure_comparison:
                lines.append("")
                lines.append("**API êµ¬ì¡° ê¸°ë°˜ ê²€ì¦:**")
                if structure_comparison.get("structure_match"):
                    lines.append("- âœ… ë°ì´í„° êµ¬ì¡°ê°€ API êµ¬ì¡°ì™€ ì¼ì¹˜í•©ë‹ˆë‹¤")
                else:
                    missing = structure_comparison.get("missing_fields", [])
                    extra = structure_comparison.get("extra_fields", [])
                    if missing:
                        lines.append(f"- âš ï¸ ëˆ„ë½ëœ í•„ë“œ ({len(missing)}ê°œ): {', '.join(missing[:5])}{'...' if len(missing) > 5 else ''}")
                    if extra:
                        lines.append(f"- â„¹ï¸ ì¶”ê°€ í•„ë“œ ({len(extra)}ê°œ): {', '.join(extra[:5])}{'...' if len(extra) > 5 else ''}")
            lines.append("")
            
            # ê²€ì¦ ì‹œê°„
            timestamp = validation_result.get("timestamp")
            if timestamp:
                lines.append(f"**ê²€ì¦ ì‹œê°„:** {timestamp}")
                lines.append("")
        
        return "\n".join(lines)
    
    def _add_markdown_analysis_section(self, lines: List[str], title: str, analysis: Dict[str, Any]):
        """Markdown ë¶„ì„ ì„¹ì…˜ ì¶”ê°€"""
        score = analysis.get('score', 0)
        grade = self._get_grade(score)
        lines.append(f"#### {title}: **{score}/100** ({grade})")
        lines.append("")
        
        if title == "ì´ë¯¸ì§€ ë¶„ì„":
            lines.append(f"- ì¸ë„¤ì¼ í’ˆì§ˆ: {analysis.get('thumbnail_quality', 'N/A')}")
            lines.append(f"- ìƒì„¸ ì´ë¯¸ì§€ ê°œìˆ˜: {analysis.get('image_count', 0)}ê°œ")
        elif title == "ìƒí’ˆ ì„¤ëª… ë¶„ì„":
            lines.append(f"- ì„¤ëª… ê¸¸ì´: {analysis.get('description_length', 0)}ì")
            lines.append(f"- êµ¬ì¡°í™” í’ˆì§ˆ: {analysis.get('structure_quality', 'N/A')}")
            keywords = analysis.get('seo_keywords', [])
            if keywords:
                lines.append(f"- SEO í‚¤ì›Œë“œ: {', '.join(keywords)}")
        elif title == "SEO ë¶„ì„":
            lines.append(f"- í‚¤ì›Œë“œ ìƒí’ˆëª… í¬í•¨: {'ì˜ˆ' if analysis.get('keywords_in_name') else 'ì•„ë‹ˆì˜¤'}")
            lines.append(f"- í‚¤ì›Œë“œ ì„¤ëª… í¬í•¨: {'ì˜ˆ' if analysis.get('keywords_in_description') else 'ì•„ë‹ˆì˜¤'}")
            lines.append(f"- ì¹´í…Œê³ ë¦¬ ì„¤ì •: {'ì˜ˆ' if analysis.get('category_set') else 'ì•„ë‹ˆì˜¤'}")
            lines.append(f"- ë¸Œëœë“œ ì„¤ì •: {'ì˜ˆ' if analysis.get('brand_set') else 'ì•„ë‹ˆì˜¤'}")
        elif title == "í˜ì´ì§€ êµ¬ì¡° ë¶„ì„":
            lines.append(f"- ì „ì²´ í´ë˜ìŠ¤ ìˆ˜: {analysis.get('total_classes', 0)}ê°œ")
            key_elements = analysis.get("key_elements_present", {})
            if key_elements:
                lines.append("- ì£¼ìš” ìš”ì†Œ ì¡´ì¬ ì—¬ë¶€:")
                for key, present in key_elements.items():
                    lines.append(f"  - {key}: {'ì˜ˆ' if present else 'ì•„ë‹ˆì˜¤'}")
        
        if analysis.get("recommendations"):
            lines.append("**ì¶”ì²œ ì‚¬í•­:**")
            for rec in analysis["recommendations"]:
                lines.append(f"- {rec}")
        lines.append("")
    
    def _add_markdown_price_analysis(self, lines: List[str], analysis: Dict[str, Any]):
        """Markdown ê°€ê²© ë¶„ì„ ì¶”ê°€ (ê°œì„ ëœ í¬ë¡¤ëŸ¬ ë°ì´í„° ë°˜ì˜)"""
        score = analysis.get('score', 0)
        grade = self._get_grade(score)
        lines.append(f"#### ê°€ê²© ë¶„ì„: **{score}/100** ({grade})")
        lines.append("")
        
        # ìœ íš¨ì„± ê²€ì¦ëœ ê°€ê²©ë§Œ í‘œì‹œ (100~1,000,000ì—” ë²”ìœ„)
        sale_price = analysis.get('sale_price')
        original_price = analysis.get('original_price')
        discount_rate = analysis.get('discount_rate', 0) or 0
        
        if sale_price and 100 <= sale_price <= 1000000:
            lines.append(f"- íŒë§¤ê°€: {sale_price:,}å††")
        else:
            lines.append(f"- íŒë§¤ê°€: N/A (ìœ íš¨í•˜ì§€ ì•Šì€ ê°’)")
        
        if original_price and 100 <= original_price <= 1000000:
            lines.append(f"- ì •ê°€: {original_price:,}å††")
            if sale_price and original_price > sale_price:
                calculated_discount = int((original_price - sale_price) / original_price * 100)
                lines.append(f"- í• ì¸ìœ¨: {calculated_discount}%")
        elif discount_rate > 0:
            lines.append(f"- í• ì¸ìœ¨: {discount_rate}%")
        
        positioning = analysis.get('positioning', '')
        if positioning:
            lines.append(f"- ê°€ê²© í¬ì§€ì…”ë‹: {positioning}")
        
        if analysis.get("recommendations"):
            lines.append("**ì¶”ì²œ ì‚¬í•­:**")
            for rec in analysis["recommendations"]:
                lines.append(f"- {rec}")
        lines.append("")
    
    def _add_markdown_review_analysis(self, lines: List[str], analysis: Dict[str, Any]):
        """Markdown ë¦¬ë·° ë¶„ì„ ì¶”ê°€ (ê°œì„ ëœ í¬ë¡¤ëŸ¬ ë°ì´í„° ë°˜ì˜)"""
        score = analysis.get('score', 0)
        grade = self._get_grade(score)
        lines.append(f"#### ë¦¬ë·° ë¶„ì„: **{score}/100** ({grade})")
        lines.append("")
        
        rating = analysis.get('rating', 0) or 0.0
        review_count = analysis.get('review_count', 0) or 0
        # fallback: reviews ë°°ì—´ ê¸¸ì´ ì‚¬ìš©
        reviews_list = analysis.get('reviews', [])
        if review_count == 0 and len(reviews_list) > 0:
            review_count = len(reviews_list)
        
        negative_ratio = analysis.get('negative_ratio', 0.0) or 0.0
        
        lines.append(f"- í‰ì : {rating:.1f}/5.0")
        if review_count > 0:
            lines.append(f"- ë¦¬ë·° ìˆ˜: {review_count:,}ê°œ")
        else:
            lines.append(f"- ë¦¬ë·° ìˆ˜: 0ê°œ (ë˜ëŠ” ì¶”ì¶œ ì‹¤íŒ¨)")
        
        if len(reviews_list) > 0:
            lines.append(f"- ì¶”ì¶œëœ ë¦¬ë·° í…ìŠ¤íŠ¸: {len(reviews_list)}ê°œ")
        
        if negative_ratio > 0:
            lines.append(f"- ë¶€ì • ë¦¬ë·° ë¹„ìœ¨: {negative_ratio:.1%}")
        
        if analysis.get("recommendations"):
            lines.append("**ì¶”ì²œ ì‚¬í•­:**")
            for rec in analysis["recommendations"]:
                lines.append(f"- {rec}")
        lines.append("")
    
    def _add_markdown_shop_info(self, lines: List[str], shop_analysis: Dict[str, Any]):
        """Markdown Shop ì •ë³´ ë¶„ì„ ì¶”ê°€"""
        shop_info = shop_analysis.get("shop_info", {})
        if shop_info:
            score = shop_info.get("score", 0)
            grade = self._get_grade(score)
            lines.append(f"#### Shop ì •ë³´ ë¶„ì„: **{score}/100** ({grade})")
            lines.append("")
        
        level_analysis = shop_analysis.get("level_analysis", {})
        if level_analysis:
            lines.append("#### Shop ë ˆë²¨ ë¶„ì„")
            lines.append("")
            lines.append(f"- í˜„ì¬ ë ˆë²¨: {level_analysis.get('current_level', 'N/A')}")
            lines.append(f"- ì •ì‚° ë¦¬ë“œíƒ€ì„: {level_analysis.get('settlement_leadtime', 15)}ì¼")
            lines.append(f"- ëª©í‘œ ë ˆë²¨: {level_analysis.get('target_level', 'N/A')}")
            lines.append("")
            
            if level_analysis.get("requirements"):
                lines.append("**ìš”êµ¬ì‚¬í•­:**")
                for req in level_analysis["requirements"]:
                    lines.append(f"- {req}")
                lines.append("")
            
            if level_analysis.get("recommendations"):
                lines.append("**ì¶”ì²œ ì‚¬í•­:**")
                for rec in level_analysis["recommendations"]:
                    lines.append(f"- {rec}")
                lines.append("")
    
    def _add_markdown_shop_specialty(self, lines: List[str], specialty: Dict[str, Any]):
        """Markdown Shop íŠ¹ìˆ˜ì„± ì¶”ê°€"""
        lines.append("#### Shop íŠ¹ìˆ˜ì„± ë¶„ì„")
        lines.append("")
        lines.append(f"- ë¸Œëœë“œ ìƒµ ì—¬ë¶€: {'ì˜ˆ' if specialty.get('is_brand_shop') else 'ì•„ë‹ˆì˜¤'}")
        if specialty.get("brand_name"):
            lines.append(f"- ë¸Œëœë“œëª…: {specialty.get('brand_name')}")
        lines.append(f"- ì œí’ˆ ë¼ì¸ì—… íŠ¹ì„±: {specialty.get('product_lineup_type', 'mixed')}")
        lines.append(f"- íƒ€ê²Ÿ ê³ ê°ì¸µ: {specialty.get('target_customer', 'general').replace('_', ' ')}")
        lines.append("")
        
        unique_features = specialty.get("unique_features", [])
        if unique_features:
            lines.append("**ë…íŠ¹í•œ íŠ¹ì§•:**")
            for feature in unique_features:
                lines.append(f"- {feature}")
            lines.append("")
        
        score = specialty.get("specialty_score", 0)
        grade = self._get_grade(score)
        lines.append(f"- íŠ¹ìˆ˜ì„± ì ìˆ˜: **{score}/100** ({grade})")
        lines.append("")
    
    def _add_markdown_customized_insights(self, lines: List[str], insights: Dict[str, Any]):
        """Markdown ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸ ì¶”ê°€"""
        lines.append("#### ë§ì¶¤í˜• ì¸ì‚¬ì´íŠ¸")
        lines.append("")
        
        if insights.get("shop_positioning"):
            lines.append(f"**Shop í¬ì§€ì…”ë‹:** {insights.get('shop_positioning')}")
            lines.append("")
        
        strengths = insights.get("strengths", [])
        if strengths:
            lines.append("**ê°•ì :**")
            for strength in strengths:
                lines.append(f"- {strength}")
            lines.append("")
        
        opportunities = insights.get("opportunities", [])
        if opportunities:
            lines.append("**ê¸°íšŒ:**")
            for opp in opportunities:
                lines.append(f"- {opp}")
            lines.append("")
        
        recommendations = insights.get("recommendations", [])
        if recommendations:
            lines.append("**ì¶”ì²œ ì‚¬í•­:**")
            for rec in recommendations:
                lines.append(f"- {rec}")
            lines.append("")
        
        advantages = insights.get("competitive_advantages", [])
        if advantages:
            lines.append("**ê²½ìŸ ìš°ìœ„:**")
            for adv in advantages:
                lines.append(f"- {adv}")
            lines.append("")